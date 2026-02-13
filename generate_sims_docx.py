from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('PROJECT TITLE', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph('SCHOOL INFORMATION MANAGEMENT SYSTEM (SIMS)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.size = Pt(16)
    p.style.font.bold = True

    # Submitted By (Placeholder)
    doc.add_paragraph('SUBMITTED BY')
    doc.add_paragraph('[Student Name]')
    doc.add_paragraph('[ID/Roll Number]')

    # Abstract
    doc.add_heading('ABSTRACT', level=1)
    doc.add_paragraph(
        "The School Information Management System (SIMS) is a comprehensive digital platform designed to streamline "
        "administrative and academic processes in educational institutions. It bridges the communication gap between "
        "administrators, teachers, students, and parents. The system facilitates the management of student records, "
        "attendance, assignments, examinations, fees, and library resources. By automating routine tasks and providing "
        "real-time access to information, SIMS enhances operational efficiency, transparency, and data accuracy, "
        "ultimately fostering a better educational environment."
    )

    # Hardware and Software Requirements
    doc.add_heading('HARDWARE AND SOFTWARE REQUIREMENTS', level=1)
    
    doc.add_heading('Hardware Requirements:', level=2)
    p = doc.add_paragraph()
    p.add_run('• Processor: ').bold = True
    p.add_run('Intel Core i5 or equivalent (Server), i3 or mobile device (Client)')
    p = doc.add_paragraph()
    p.add_run('• RAM: ').bold = True
    p.add_run('8 GB minimum (Server), 4 GB (Client)')
    p = doc.add_paragraph()
    p.add_run('• Storage: ').bold = True
    p.add_run('SSD recommended for Database Server')
    p = doc.add_paragraph()
    p.add_run('• Network: ').bold = True
    p.add_run('Stable Broadband/Internet connection')

    doc.add_heading('Software Requirements:', level=2)
    p = doc.add_paragraph()
    p.add_run('• Frontend: ').bold = True
    p.add_run('React.js, Vite, Tailwind CSS/Bootstrap')
    p = doc.add_paragraph()
    p.add_run('• Backend: ').bold = True
    p.add_run('Python, FastAPI')
    p = doc.add_paragraph()
    p.add_run('• Database: ').bold = True
    p.add_run('PostgreSQL')
    p = doc.add_paragraph()
    p.add_run('• ORM: ').bold = True
    p.add_run('SQLAlchemy, Alembic (Migrations)')
    p = doc.add_paragraph()
    p.add_run('• Tools: ').bold = True
    p.add_run('Visual Studio Code, Docker, Git, Postman')

    # Existing vs Proposed System
    doc.add_heading('EXISTING SYSTEM & PROPOSED SYSTEM', level=1)
    
    doc.add_heading('Existing System', level=2)
    doc.add_paragraph("• Manual Record Keeping: Physical registers for attendance and marks.")
    doc.add_paragraph("• Disconnected Tools: Separate excel sheets for fees, results, and student details.")
    doc.add_paragraph("• Communication Gaps: Reliance on physical notice boards or paper diaries.")
    doc.add_paragraph("• Data Redundancy: Same information repeated across different manual logs.")

    doc.add_heading('Proposed System', level=2)
    doc.add_paragraph("• Centralized Database: Single source of truth for all school data.")
    doc.add_paragraph("• Role-Based Access: Secure dashboards for Admins, Teachers, Students, and Parents.")
    doc.add_paragraph("• Automation: Automatic fee calculations, result generation, and attendance tracking.")
    doc.add_paragraph("• Remote Access: Web-based platform accessible from anywhere.")
    doc.add_paragraph("• Digital Library & Assets: Management of book lending and school inventory.")

    # Literature Survey
    doc.add_heading('LITERATURE SURVEY', level=1)
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        "A study of existing School Management Systems reveals a shift from legacy desktop applications to "
        "cloud-based solutions. However, many existing solutions are either too expensive for small-to-medium "
        "schools or lack user-friendly interfaces."
    )
    
    doc.add_heading('2. Study of Existing Systems', level=2)
    doc.add_paragraph(
        "• 'Moodle' (Open Source): Primarily an LMS, lacks comprehensive administrative features like fee management.\n"
        "• 'Blackboard': Powerful but expensive and complex for smaller institutions.\n"
        "• Custom Excel/Access Solutions: Prone to data corruption and lack multi-user concurrency."
    )

    doc.add_heading('3. Research Gap', level=2)
    doc.add_paragraph(
        "There is a need for a lightweight, modern, and cost-effective solution that combines Learning Management (LMS) "
        "features with Administrative (ERP) capabilities. SIMS addresses this by using a modern tech stack (FastAPI/React) "
        "to ensure performance and scalability without the bloat of enterprise legacy systems."
    )

    # Research Methodology
    doc.add_heading('RESEARCH METHODOLOGY', level=1)
    
    doc.add_heading('1. Research Approach', level=2)
    doc.add_paragraph("Applied Research targeting the operational inefficiencies in educational institutions.")

    doc.add_heading('2. System Architecture', level=2)
    doc.add_paragraph(
        "The system follows a Client-Server Architecture (RESTful API).\n"
        "• Client: React SPA (Single Page Application) consuming JSON APIs.\n"
        "• Server: FastAPI (Python) handling business logic and ORM mapping.\n"
        "• Database: PostgreSQL for relational data storage."
    )

    doc.add_heading('3. Module Description', level=2)
    modules = [
        ("Authentication", "Handles Login, Registration, Password Reset using JWT."),
        ("Admin Dashboard", "User management (Teachers/Students), Global settings."),
        ("Student Management", "Profile management, Class assignment."),
        ("Teacher Management", "Staff profiles, Subject allocation."),
        ("Academic Module", "Classes, Sections, Subjects, Timetable management."),
        ("Assignment Module", "Creation, Submission, and Grading of assignments."),
        ("Examination & Marks", "Exam scheduling, Mark entry, Report card generation."),
        ("Attendance", "Daily attendance tracking for students and staff."),
        ("Fees Management", "Fee structure creation, Payment tracking, Due alerts."),
        ("Library", "Book catalog, Issue/Return tracking."),
        ("Communication", "Notifications, Events, Parent-Teacher updates.")
    ]
    for name, desc in modules:
        p = doc.add_paragraph()
        p.add_run(f"• {name}: ").bold = True
        p.add_run(desc)

    doc.add_heading('4. Process Flow', level=2)
    doc.add_paragraph("1. User logs in with credentials.")
    doc.add_paragraph("2. System validates Role (Admin/Teacher/Student/Parent).")
    doc.add_paragraph("3. User is redirected to their specific Dashboard.")
    doc.add_paragraph("4. User performs actions (e.g., Mark Attendance, Pay Fees).")
    doc.add_paragraph("5. Backend validates request and updates Database.")
    doc.add_paragraph("6. Success/Error response displayed to User.")

    # System Diagrams
    doc.add_heading('SYSTEM DIAGRAMS', level=1)
    
    doc.add_heading('Architecture Diagram', level=2)
    doc.add_paragraph("[Diagram Description: A 3-tier architecture showing the React Frontend interacting with the FastAPI Backend via HTTP/REST, which in turn communicates with the PostgreSQL Database using SQLAlchemy.]")
    
    doc.add_heading('Class Diagram (Key Entities)', level=2)
    doc.add_paragraph(
        "Key Classes/Entities:\n"
        "- User (Base for Admin, Teacher, Student, Parent)\n"
        "- ClassRoom (has many Students)\n"
        "- Subject (taught by Teacher)\n"
        "- Assignment (linked to Subject and Class)\n"
        "- Submission (linked to Assignment and Student)\n"
        "- Attendance (linked to Student and Date)\n"
        "- Fee (linked to Student)"
    )

    doc.add_heading('Use Case Diagram', level=2)
    doc.add_paragraph(
        "Actors:\n"
        "- Admin: Add Users, Manage Fees, Configure System.\n"
        "- Teacher: Take Attendance, Upload Assignments, Enter Marks.\n"
        "- Student: View Timetable, Submit Assignments, View Marks.\n"
        "- Parent: View Child's Attendance, Pay Fees, View Report Cards."
    )

    doc.add_heading('ER Diagram', level=2)
    doc.add_paragraph(
        "[Diagram Description: Relational model showing One-to-Many relationships between Class and Students, "
        "Many-to-Many between Teachers and Subjects, One-to-Many between Students and Fee Records, etc.]"
    )

    doc.save('PROJECT_PHASE_2.docx')
    print("Document created successfully: PROJECT_PHASE_2.docx")

if __name__ == "__main__":
    create_document()
